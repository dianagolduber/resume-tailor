#!/usr/bin/env python3
"""
Resume Tailor — automatically tailors a resume to a job description using Claude.
Usage:
    python tailor.py                              # paste job description interactively
    python tailor.py --url "https://..."          # scrape job from a URL
    python tailor.py --jd job.txt                 # load job description from a file
    python tailor.py --url "https://..." --company "Stripe" --role "PMM"

Required env vars:
    ANTHROPIC_API_KEY   your Claude API key
    BASE_RESUME_PATH    path to your base resume PDF (or pass --resume flag)

Optional env vars:
    OUTPUT_DIR          where to save tailored resumes (default: ./output)
"""

import argparse
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import anthropic
import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

BASE_RESUME_PATH = Path(os.environ.get("BASE_RESUME_PATH", "resume.pdf"))
OUTPUT_DIR = Path(os.environ.get("OUTPUT_DIR", Path(__file__).parent / "output"))

SYSTEM_PROMPT = """You are an expert resume writer and career coach specializing in marketing and product marketing roles.
Your job is to tailor a resume to a specific job description while keeping all facts 100% accurate — never invent experience, metrics, or skills.

When tailoring:
- Reorder bullet points so the most relevant experience appears first
- Mirror keywords and phrases from the job description naturally (for ATS systems)
- Strengthen weak bullet points by making them more specific and impact-focused
- Adjust the summary/intro paragraph to speak directly to the role
- Remove or de-emphasize experience that isn't relevant
- Keep formatting clean and consistent
- Do NOT add fake metrics or experience that isn't in the original

Output the full tailored resume as plain text, preserving the structure (name, contact, summary, experience, skills, education).
At the end, add a section called "## TAILORING NOTES" explaining the top 3-5 changes you made and why."""


def read_resume_pdf(path: Path) -> str:
    with pdfplumber.open(path) as pdf:
        pages = [page.extract_text() for page in pdf.pages if page.extract_text()]
    return "\n\n".join(pages)


def scrape_job_url(url: str) -> tuple[str, str, str]:
    """Scrape a job posting URL. Returns (text, company, role)."""
    print(f"Fetching {url} ...")
    headers = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"}
    try:
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
    except requests.RequestException as e:
        print(f"Error fetching URL: {e}")
        sys.exit(1)

    soup = BeautifulSoup(resp.text, "html.parser")

    # Remove nav, footer, scripts, styles
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    # Try common job posting containers first
    job_section = (
        soup.find(attrs={"data-testid": re.compile(r"job|description", re.I)})
        or soup.find(class_=re.compile(r"job[_-]?description|posting|jd|description", re.I))
        or soup.find("main")
        or soup.find("article")
        or soup.body
    )

    text = job_section.get_text(separator="\n") if job_section else soup.get_text(separator="\n")

    # Clean up whitespace
    lines = [l.strip() for l in text.splitlines()]
    text = "\n".join(l for l in lines if l)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Try to auto-detect company and role from page title
    title = soup.title.string if soup.title else ""
    company, role = "", ""
    # Common patterns: "Role at Company", "Role | Company", "Company - Role"
    for pattern in [r"^(.+?)\s+at\s+(.+?)(?:\s*[\|\-]|$)", r"^(.+?)\s*[\|–\-]\s*(.+?)$"]:
        m = re.match(pattern, title.strip())
        if m:
            role, company = m.group(1).strip(), m.group(2).strip()
            break

    return text, company, role


def get_job_description(args) -> tuple[str, str, str]:
    """Returns (job_description_text, company, role)."""
    company = args.company or ""
    role = args.role or ""

    if args.url:
        text, scraped_company, scraped_role = scrape_job_url(args.url)
        # Command-line flags override auto-detected values
        company = company or scraped_company
        role = role or scraped_role
        print(f"Scraped {len(text.split())} words from URL.")
        return text, company, role

    if args.jd:
        jd_path = Path(args.jd)
        if not jd_path.exists():
            print(f"Error: file not found: {args.jd}")
            sys.exit(1)
        return jd_path.read_text(), company, role

    print("\nPaste the job description below.")
    print("When done, type END on its own line and press Enter:\n")
    lines = []
    while True:
        try:
            line = input()
            if line.strip() == "END":
                break
            lines.append(line)
        except EOFError:
            break
    return "\n".join(lines), company, role


def save_outputs(content: str, company: str, role: str) -> tuple[Path, Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    slug = f"{company.replace(' ', '_')}_{role.replace(' ', '_')}_{timestamp}" if company or role else timestamp

    # Save plain text
    txt_path = OUTPUT_DIR / f"resume_{slug}.txt"
    txt_path.write_text(content)

    # Save as .docx
    docx_path = OUTPUT_DIR / f"resume_{slug}.docx"
    doc = Document()

    # Basic styling
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in content.split("\n"):
        if line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        else:
            p = doc.add_paragraph(line)

    doc.save(docx_path)
    return txt_path, docx_path


def main():
    parser = argparse.ArgumentParser(description="Tailor a resume to a job description")
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--url", help="URL of the job posting to scrape")
    group.add_argument("--jd", help="Path to a text file containing the job description")
    parser.add_argument("--resume", help="Path to your base resume PDF (overrides BASE_RESUME_PATH env var)")
    parser.add_argument("--company", default="", help="Company name (overrides auto-detected, used in filename)")
    parser.add_argument("--role", default="", help="Role/title (overrides auto-detected, used in filename)")
    args = parser.parse_args()

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("Error: ANTHROPIC_API_KEY environment variable not set.")
        print("Set it with: export ANTHROPIC_API_KEY=your_key_here")
        sys.exit(1)

    resume_path = Path(args.resume) if args.resume else BASE_RESUME_PATH
    if not resume_path.exists():
        print(f"Error: resume not found at {resume_path}")
        print("Set BASE_RESUME_PATH env var or use --resume flag.")
        sys.exit(1)

    print("Reading base resume...")
    resume_text = read_resume_pdf(resume_path)

    print("Loading job description...")
    job_description, company, role = get_job_description(args)

    if not job_description.strip():
        print("Error: job description is empty.")
        sys.exit(1)

    company_label = f" at {company}" if company else ""
    role_label = f" for {role}" if role else ""
    print(f"\nTailoring resume{role_label}{company_label}...\n")
    print("-" * 60)

    client = anthropic.Anthropic(api_key=api_key)

    user_message = f"""Here is my current resume:

<resume>
{resume_text}
</resume>

Here is the job description I'm applying to:

<job_description>
{job_description}
</job_description>

Please tailor my resume for this role. Remember: only use facts from my actual resume — do not invent anything."""

    full_response = ""
    with client.messages.stream(
        model="claude-opus-4-6",
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        for text in stream.text_stream:
            print(text, end="", flush=True)
            full_response += text

    print("\n" + "-" * 60)

    txt_path, docx_path = save_outputs(full_response, company, role)
    print(f"\nSaved:")
    print(f"  Text: {txt_path}")
    print(f"  Word: {docx_path}")

    # Auto-log to tracker
    try:
        import sqlite3
        from tracker import get_db, STATUSES
        conn = get_db()
        now = datetime.now().strftime("%Y-%m-%d")
        url = args.url or ""
        conn.execute(
            """INSERT INTO applications (company, role, url, status, resume_path, notes, date_applied, updated_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (company or "Unknown", role or "Unknown", url, "Applied", str(docx_path), "", now, now)
        )
        conn.commit()
        print(f"\n\033[32mLogged to tracker.\033[0m  Run: python tracker.py list")
    except Exception as e:
        print(f"\n(Tracker log failed: {e})")


if __name__ == "__main__":
    main()
